"""Document text extraction for the AI assistant's file-upload feature (Phase 4).

Deliberately separate from services/storage_service.py's validate_upload_file /
ALLOWED_DOCUMENT_EXTENSIONS — that shared function is used by every OTHER upload route
in the app (resumes, employer documents, announcements, job fair materials), none of
which should suddenly start accepting DOCX/TXT. This module is scoped to chat uploads only.

Nothing here ever touches disk or Supabase Storage — file bytes live only in the request
that uploaded them (see blueprints/assistant.py's upload_document route) and are handed
to pypdf/python-docx via an in-memory BytesIO. There is nothing to delete because nothing
durable is ever created.

PDF is read with pypdf (native text-layer extraction — free, instant, no dependency on
Google Vision being configured/billed). This only reads real text layers: a scanned/
image-only PDF with no selectable text is reported as an extraction failure rather than
silently producing near-empty content the model would have to guess at, and rather than
routing through the paid OCR pipeline resumes already use (services/ocr_service.py) —
that tradeoff is deliberate, see the Phase 4 plan.
"""

import os

import docx
from pypdf import PdfReader

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}
MAX_SIZE_BYTES = 5 * 1024 * 1024  # 5MB — same ceiling every other upload type in this app uses

# The real token-cost control isn't file size — a 5MB DOCX can still be hundreds of pages
# of raw text, so extracted text is truncated after the fact regardless of source length.
MAX_EXTRACTED_TEXT_CHARS = 8000
MAX_PDF_PAGES = 10
MIN_EXTRACTED_TEXT_CHARS = 50  # below this, treat a "PDF" as having no real text layer

_MAGIC_BYTES = {
    "pdf": b"%PDF-",
    "docx": b"PK\x03\x04",  # docx is a zip archive; a mismatch here is caught for real by
}                            # the actual parse attempt below, which is the authoritative check


def sanitize_filename(filename: str) -> str:
    """Display-only — this filename is never used as a storage path (nothing is persisted),
    only echoed in error messages and dropped into the LLM prompt. No sanitizer exists
    elsewhere in this codebase since every other upload route sidesteps the problem by
    writing to a uuid-named path instead of the original filename."""
    name = os.path.basename(filename or "upload").strip()
    name = "".join(ch for ch in name if ch.isprintable())
    return name[:150] or "upload"


def validate(file_bytes: bytes, filename: str) -> str | None:
    """Returns an error message if the file is invalid, else None."""
    if not file_bytes:
        return "The uploaded file is empty."
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return "Only PDF, DOCX, and TXT files are allowed."
    if len(file_bytes) > MAX_SIZE_BYTES:
        return "File is too large. Maximum size is 5MB."
    magic = _MAGIC_BYTES.get(ext)
    if magic and not file_bytes.startswith(magic):
        return f"This file doesn't look like a valid {ext.upper()} file."
    return None


def _extract_pdf(file_bytes: bytes) -> str:
    import io

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = reader.pages[:MAX_PDF_PAGES]
    return "\n".join(page.extract_text() or "" for page in pages).strip()


def _extract_docx(file_bytes: bytes) -> str:
    import io

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(part for part in parts if part.strip()).strip()


def extract_text(file_bytes: bytes, filename: str) -> dict:
    """Returns {"text": str|None, "mode": "real"|"error", "detail": str|None} — mirrors
    ocr_service.extract_text_from_resume()'s contract so callers check `mode` the same way."""
    ext = filename.rsplit(".", 1)[-1].lower()

    try:
        if ext == "pdf":
            text = _extract_pdf(file_bytes)
            if len(text) < MIN_EXTRACTED_TEXT_CHARS:
                return {
                    "text": None, "mode": "error",
                    "detail": "This PDF doesn't appear to contain selectable text (it may be a "
                              "scanned image). Please upload a text-based PDF, DOCX, or TXT file instead.",
                }
        elif ext == "docx":
            text = _extract_docx(file_bytes)
            if len(text) < MIN_EXTRACTED_TEXT_CHARS:
                return {"text": None, "mode": "error", "detail": "This DOCX file appears to have no readable text content."}
        elif ext == "txt":
            try:
                text = file_bytes.decode("utf-8").strip()
            except UnicodeDecodeError:
                return {"text": None, "mode": "error", "detail": "This file doesn't appear to be valid UTF-8 text."}
            if not text:
                return {"text": None, "mode": "error", "detail": "This text file appears to be empty."}
        else:
            return {"text": None, "mode": "error", "detail": "Unsupported file type."}
    except Exception as exc:  # noqa: BLE001
        detail = (
            "Couldn't read this PDF — it may be corrupted or password-protected."
            if ext == "pdf" else
            "Couldn't read this DOCX file — it may be corrupted."
        )
        return {"text": None, "mode": "error", "detail": detail}

    if len(text) > MAX_EXTRACTED_TEXT_CHARS:
        text = text[:MAX_EXTRACTED_TEXT_CHARS].rsplit(" ", 1)[0] + "…"
    return {"text": text, "mode": "real", "detail": None}
