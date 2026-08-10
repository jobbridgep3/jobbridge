"""Resume document parsing via Gemini, for Profile Management's resume upload/extraction
feature (POST /api/profile/resume). Replaces the old Google Cloud Vision OCR + regex/
spaCy field-mapper pipeline for THIS feature only — Vision itself is still used
elsewhere (SPES/DILP/OWWA program-document text extraction, services/ocr_service.py),
which is unrelated and untouched by this module.

This is document understanding, not raw OCR: Gemini reads the resume and maps
explicitly stated information directly onto the structured fields below, via
response_mime_type=application/json + response_schema so the output is always valid
JSON — no manual regex/text parsing on our end. Gemini is instructed to leave a field
blank/empty rather than infer, assume, calculate, or guess a value that isn't
explicitly stated in the document; _normalize() below is a second, independent
backend-side check that enforces the expected shape regardless of what Gemini returns
(never trusts the model's output as already-safe).

Plain HTTPS via `requests` — no vendor SDK — matching services/assistant_service.py's
existing Groq integration convention (see that module's docstring for the rationale).
The API key comes from GEMINI_API_KEY (env only; never sent to the browser).

eventlet: the app runs on a single eventlet worker (see render.yaml startCommand), and
`requests`/urllib3 socket I/O is not covered by eventlet's monkey_patch() in a way
that's safe to call directly from a greenlet. The HTTP call is offloaded via
eventlet.tpool.execute(), exactly as assistant_service.py and (formerly) ocr_service.py
do. tpool dispatches to real OS threads that do NOT inherit the Flask app context, so
all current_app.config reads happen before entering tpool, never inside it.
"""

import base64
import json
import logging

import requests

logger = logging.getLogger(__name__)

GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"

# Connect / read timeouts. The read timeout must stay well inside gunicorn's 90s worker
# timeout so a slow upstream surfaces as our own clean error rather than a killed worker.
_TIMEOUT = (10, 45)

ALLOWED_EXTENSIONS = {"pdf", "docx", "jpg", "jpeg", "png"}
# Same 5MB ceiling every other upload type in this app uses (storage_service.py,
# document_extraction.py) — deliberately not shared code, since this validator accepts
# DOCX and those don't (see document_extraction.py's docstring for why that's scoped
# per-feature rather than centralized).
MAX_SIZE_BYTES = 5 * 1024 * 1024

_MAGIC_BYTES = {
    "pdf": b"%PDF-",
    "docx": b"PK\x03\x04",  # docx is a zip archive; a mismatch here is caught for real
    "jpg": b"\xff\xd8\xff",  # by the actual Gemini/python-docx read attempt below, which
    "jpeg": b"\xff\xd8\xff",  # is the authoritative check
    "png": b"\x89PNG\r\n\x1a\n",
}

_MIME_TYPES = {"pdf": "application/pdf", "jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png"}

_ERR_UNCONFIGURED = "Resume extraction is not configured on this server."
_ERR_TIMEOUT = "Resume extraction took too long. Please try again."
_ERR_UNAVAILABLE = "Resume extraction is temporarily unavailable. Please try again in a moment."
_ERR_MALFORMED = "Could not read the extracted data from this resume. Please fill in your details manually."

_PERSONAL_FIELDS = [
    "full_name", "email", "contact_number", "date_of_birth", "gender", "civil_status",
    "nationality", "barangay", "municipality", "province", "region", "zip_code",
]
_EDUCATION_FIELDS = ["school_name", "level", "course", "year", "honors"]
_EMPLOYMENT_FIELDS = [
    "employment_status", "employment_type_preferred", "preferred_job_position",
    "preferred_industry", "preferred_work_location", "expected_salary",
]
_WORK_EXPERIENCE_FIELDS = ["company", "position", "start_date", "end_date", "description"]
_SKILLS_LIST_FIELDS = ["technical_skills", "soft_skills", "languages_spoken", "certifications"]

_PROMPT = (
    "You are a resume document-parsing assistant. Read this resume and extract ONLY "
    "information that is explicitly and literally stated in the document, mapped onto "
    "the provided JSON schema. Do not fabricate, assume, infer, calculate, or derive "
    "any value that is not explicitly written in the resume (for example: do not guess "
    "a birth year from an age, do not infer a skill from a job title, do not calculate "
    "years of experience). If a field is not explicitly stated in the resume, return an "
    "empty string for that field, or an empty array for list fields. Extract every "
    "education entry and every work experience entry present in the document, in the "
    "order they appear."
)

_SCHEMA = {
    "type": "OBJECT",
    "properties": {
        "personal_information": {
            "type": "OBJECT",
            "properties": {field: {"type": "STRING"} for field in _PERSONAL_FIELDS},
            "required": _PERSONAL_FIELDS,
        },
        "educational_background": {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {field: {"type": "STRING"} for field in _EDUCATION_FIELDS},
                "required": _EDUCATION_FIELDS,
            },
        },
        "employment_information": {
            "type": "OBJECT",
            "properties": {
                **{field: {"type": "STRING"} for field in _EMPLOYMENT_FIELDS},
                "work_experience": {
                    "type": "ARRAY",
                    "items": {
                        "type": "OBJECT",
                        "properties": {field: {"type": "STRING"} for field in _WORK_EXPERIENCE_FIELDS},
                        "required": _WORK_EXPERIENCE_FIELDS,
                    },
                },
            },
            "required": [*_EMPLOYMENT_FIELDS, "work_experience"],
        },
        "skills": {
            "type": "OBJECT",
            "properties": {field: {"type": "ARRAY", "items": {"type": "STRING"}} for field in _SKILLS_LIST_FIELDS},
            "required": _SKILLS_LIST_FIELDS,
        },
    },
    "required": ["personal_information", "educational_background", "employment_information", "skills"],
}


def is_gemini_configured() -> bool:
    from flask import current_app

    return bool(current_app.config.get("GEMINI_API_KEY"))


def validate_resume_file(file_bytes: bytes, filename: str) -> str | None:
    """Returns an error message if the file is invalid, else None."""
    if not file_bytes:
        return "The uploaded file is empty."
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return "Only PDF, DOCX, JPG, and PNG files are allowed."
    if len(file_bytes) > MAX_SIZE_BYTES:
        return "File is too large. Maximum size is 5MB."
    magic = _MAGIC_BYTES.get(ext)
    if magic and not file_bytes.startswith(magic):
        return f"This file doesn't look like a valid {ext.upper()} file."
    return None


def _extract_docx_text(file_bytes: bytes) -> str:
    """Gemini has no native DOCX input type, so DOCX resumes are text-extracted first
    and sent as a text part instead of inline file bytes."""
    import io

    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(part for part in parts if part.strip()).strip()


def _build_parts(file_bytes: bytes, filename: str) -> list:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "docx":
        text = _extract_docx_text(file_bytes)
        if not text:
            raise ValueError("This DOCX file appears to have no readable text content.")
        return [{"text": f"{_PROMPT}\n\nResume content:\n\n{text}"}]

    return [
        {"text": _PROMPT},
        {"inline_data": {"mime_type": _MIME_TYPES[ext], "data": base64.b64encode(file_bytes).decode("ascii")}},
    ]


def _coerce_str(value) -> str:
    return value.strip() if isinstance(value, str) else ""


def _coerce_str_list(value) -> list:
    if not isinstance(value, list):
        return []
    return [item.strip() for item in value if isinstance(item, str) and item.strip()]


def _normalize(data: dict) -> dict:
    """Backend-side structural validation/normalization — never trusts Gemini's JSON as
    already-safe. Coerces missing/malformed fields to "" or [] rather than raising, so a
    partially-malformed response still degrades to a mostly-empty (never fabricated)
    result instead of a hard failure."""
    if not isinstance(data, dict):
        raise ValueError("Response was not a JSON object.")

    personal_raw = data.get("personal_information")
    personal = {
        field: _coerce_str((personal_raw or {}).get(field)) for field in _PERSONAL_FIELDS
    }

    education_raw = data.get("educational_background")
    education = [
        {field: _coerce_str(entry.get(field)) for field in _EDUCATION_FIELDS}
        for entry in (education_raw if isinstance(education_raw, list) else [])
        if isinstance(entry, dict)
    ]

    employment_raw = data.get("employment_information") or {}
    work_experience_raw = employment_raw.get("work_experience") if isinstance(employment_raw, dict) else None
    employment = {
        **{field: _coerce_str(employment_raw.get(field)) for field in _EMPLOYMENT_FIELDS},
        "work_experience": [
            {field: _coerce_str(entry.get(field)) for field in _WORK_EXPERIENCE_FIELDS}
            for entry in (work_experience_raw if isinstance(work_experience_raw, list) else [])
            if isinstance(entry, dict)
        ],
    }

    skills_raw = data.get("skills") or {}
    skills = {field: _coerce_str_list(skills_raw.get(field)) for field in _SKILLS_LIST_FIELDS}

    return {
        "personal_information": personal,
        "educational_background": education,
        "employment_information": employment,
        "skills": skills,
    }


def _call_gemini(api_key: str, model: str, parts: list) -> dict:
    """Runs entirely inside eventlet.tpool's real OS thread — no Flask context access.
    `parts` are plain data built by the caller before entering tpool — safe to cross the
    thread boundary. Returns the parsed (but not yet normalized) JSON dict, or raises."""
    resp = requests.post(
        GEMINI_API_URL.format(model=model),
        headers={"x-goog-api-key": api_key, "Content-Type": "application/json"},
        json={
            "contents": [{"parts": parts}],
            "generationConfig": {"response_mime_type": "application/json", "response_schema": _SCHEMA},
        },
        timeout=_TIMEOUT,
    )

    if resp.status_code >= 300:
        # resp.text, not an exception chain — Gemini returns a JSON error body
        # explaining bad key / unknown model / quota, which is what you need in logs.
        raise RuntimeError(f"Gemini returned {resp.status_code}: {resp.text[:500]}")

    body = resp.json()
    candidates = body.get("candidates") or []
    if not candidates:
        raise RuntimeError(f"Gemini returned no candidates: {str(body)[:500]}")

    text_parts = (candidates[0].get("content") or {}).get("parts") or []
    text = "".join(part.get("text", "") for part in text_parts).strip()
    if not text:
        raise RuntimeError(f"Gemini returned an empty response: {str(body)[:500]}")

    return json.loads(text)


def extract_resume(file_bytes: bytes, filename: str) -> dict:
    """Returns {"extracted": dict|None, "mode": "real"|"error", "detail": str|None}.

    Callers must check `mode` rather than assume success — any failure (unconfigured,
    auth, quota, network, timeout, malformed response) is reported as mode "error" with
    no data, never fabricated placeholder data merged into a real user's profile as if
    it were a genuine extraction."""
    if not is_gemini_configured():
        logger.warning("GEMINI_API_KEY is not configured — cannot process %s", filename)
        return {"extracted": None, "mode": "error", "detail": _ERR_UNCONFIGURED}

    try:
        parts = _build_parts(file_bytes, filename)
    except ValueError as exc:
        return {"extracted": None, "mode": "error", "detail": str(exc)}

    # Read config BEFORE entering tpool — the offloaded thread has no app context.
    from flask import current_app

    api_key = current_app.config["GEMINI_API_KEY"]
    model = current_app.config["GEMINI_MODEL"]

    import eventlet.tpool

    try:
        raw = eventlet.tpool.execute(_call_gemini, api_key, model, parts)
        extracted = _normalize(raw)
        return {"extracted": extracted, "mode": "real", "detail": None}
    except requests.Timeout as exc:
        logger.error("Gemini request timed out for %s: %s", filename, exc)
        return {"extracted": None, "mode": "error", "detail": _ERR_TIMEOUT}
    except (ValueError, KeyError, TypeError) as exc:
        logger.error("Gemini returned a malformed response for %s: %s", filename, exc)
        return {"extracted": None, "mode": "error", "detail": _ERR_MALFORMED}
    except Exception as exc:  # noqa: BLE001
        logger.error("Gemini request failed for %s: %s", filename, exc)
        return {"extracted": None, "mode": "error", "detail": _ERR_UNAVAILABLE}
