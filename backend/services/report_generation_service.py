"""On-demand report file generation for the AI assistant's "generate me a list" feature.

Every report below is a small, fixed, whitelisted query returning (title, columns,
rows) — never freeform AI-written SQL, and never a query dynamically constructed from
message text. Scoping mirrors assistant_retrieval.py exactly: filtered by an id derived
from the caller's JWT identity, never anything client-controlled. `REPORT_CATALOG`'s
`roles` set is the actual security boundary — generate_report_file() refuses anything
not in the caller's own role's whitelist before a single query runs.

Rendering reuses the app's existing export infrastructure wherever it already exists
(excel_service.build_excel_report, pdf_service.generate_table_report — the same
functions every staff/admin/employer export route already calls) rather than
reinventing file generation. DOCX and TXT have no prior write precedent in this
codebase (Phase 4's python-docx use was read-only) — added here as small, self-contained
writers.
"""

import html
import io

from models.application import APPLICATION_STATUS_LABELS, Application
from models.audit import AuditTrail
from models.employer import EmployerCompany
from models.employment import EMPLOYMENT_STATUS_LABELS, EmploymentRecord
from models.jobseeker import JobseekerProfile
from models.program import ProgramApplication
from models.referral import ReferralLetter
from models.vacancy import Vacancy
from utils.timezone import now_manila


def _esc(value) -> str:
    """HTML-escapes a single value — only needed on the PDF path (generate_table_report
    interpolates cells directly into an HTML table), not Excel/DOCX/TXT."""
    return html.escape(str(value)) if value is not None else ""


def _date(value) -> str:
    return value.strftime("%Y-%m-%d") if value else ""


# ---------- Jobseeker ----------

def _my_applications(identity, date_from=None):
    profile = JobseekerProfile.query.filter_by(user_id=identity).first()
    columns = ["Vacancy", "Employer", "Status", "Date Applied"]
    if not profile:
        return "My Applications", columns, []
    query = Application.query.filter_by(jobseeker_profile_id=profile.id)
    if date_from:
        query = query.filter(Application.created_at >= date_from)
    apps = query.order_by(Application.created_at.desc()).all()
    rows = [
        [a.vacancy.title, a.vacancy.employer_company.company_name if a.vacancy.employer_company else "",
         APPLICATION_STATUS_LABELS.get(a.status, a.status), _date(a.created_at)]
        for a in apps
    ]
    return "My Applications", columns, rows


def _my_referrals(identity, date_from=None):
    profile = JobseekerProfile.query.filter_by(user_id=identity).first()
    columns = ["Vacancy", "Status", "Date"]
    if not profile:
        return "My Referral Letters", columns, []
    query = ReferralLetter.query.filter_by(jobseeker_profile_id=profile.id)
    if date_from:
        query = query.filter(ReferralLetter.created_at >= date_from)
    letters = query.order_by(ReferralLetter.created_at.desc()).all()
    rows = [[letter.vacancy.title if letter.vacancy else "General", letter.status, _date(letter.created_at)] for letter in letters]
    return "My Referral Letters", columns, rows


def _my_employment(identity, date_from=None):
    profile = JobseekerProfile.query.filter_by(user_id=identity).first()
    columns = ["Employer", "Position", "Status", "Start Date"]
    if not profile:
        return "My Employment History", columns, []
    query = EmploymentRecord.query.filter_by(jobseeker_profile_id=profile.id)
    if date_from:
        query = query.filter(EmploymentRecord.start_date >= date_from)
    records = query.order_by(EmploymentRecord.start_date.desc()).all()
    rows = [
        [r.employer_company.company_name if r.employer_company else "", r.position,
         EMPLOYMENT_STATUS_LABELS.get(r.status, r.status), _date(r.start_date)]
        for r in records
    ]
    return "My Employment History", columns, rows


# ---------- Employer ----------

def _my_vacancies(identity, date_from=None):
    company = EmployerCompany.query.filter_by(user_id=identity).first()
    columns = ["Title", "Status", "Slots", "Posted"]
    if not company:
        return "My Job Postings", columns, []
    query = Vacancy.query.filter_by(employer_company_id=company.id, is_template=False).filter(Vacancy.deleted_at.is_(None))
    if date_from:
        query = query.filter(Vacancy.created_at >= date_from)
    vacancies = query.order_by(Vacancy.created_at.desc()).all()
    rows = [[v.title, v.status, v.num_slots, _date(v.created_at)] for v in vacancies]
    return "My Job Postings", columns, rows


def _my_applicants(identity, date_from=None):
    company = EmployerCompany.query.filter_by(user_id=identity).first()
    columns = ["Applicant", "Position", "Status", "Match %", "Date Applied"]
    if not company:
        return "My Applicants", columns, []
    query = Application.query.join(Vacancy).filter(Vacancy.employer_company_id == company.id)
    if date_from:
        query = query.filter(Application.created_at >= date_from)
    apps = query.order_by(Application.created_at.desc()).all()
    rows = [
        [a.jobseeker_profile.full_name, a.vacancy.title, APPLICATION_STATUS_LABELS.get(a.status, a.status),
         f"{a.match_score:.0f}%" if a.match_score is not None else "", _date(a.created_at)]
        for a in apps
    ]
    return "My Applicants", columns, rows


def _my_hires(identity, date_from=None):
    company = EmployerCompany.query.filter_by(user_id=identity).first()
    columns = ["Employee", "Position", "Status", "Start Date"]
    if not company:
        return "My Hires", columns, []
    query = EmploymentRecord.query.filter_by(employer_company_id=company.id)
    if date_from:
        query = query.filter(EmploymentRecord.start_date >= date_from)
    records = query.order_by(EmploymentRecord.start_date.desc()).all()
    rows = [
        [r.jobseeker_profile.full_name if r.jobseeker_profile else "", r.position,
         EMPLOYMENT_STATUS_LABELS.get(r.status, r.status), _date(r.start_date)]
        for r in records
    ]
    return "My Hires", columns, rows


# ---------- PESO Staff / Admin ----------

def _staff_jobseekers(identity, date_from=None):
    query = JobseekerProfile.query
    if date_from:
        query = query.filter(JobseekerProfile.created_at >= date_from)
    profiles = query.order_by(JobseekerProfile.created_at.desc()).limit(2000).all()
    rows = [
        [p.full_name, p.municipality or "", p.employment_status or "",
         "Verified" if p.is_verified_by_staff else "Unverified", _date(p.created_at)]
        for p in profiles
    ]
    return "Jobseekers List", ["Name", "Municipality", "Employment Status", "Verification", "Registered"], rows


def _staff_employers(identity, date_from=None):
    query = EmployerCompany.query
    if date_from:
        query = query.filter(EmployerCompany.created_at >= date_from)
    companies = query.order_by(EmployerCompany.created_at.desc()).limit(2000).all()
    rows = [[c.company_name or "", c.industry or "", c.accreditation_status, _date(c.created_at)] for c in companies]
    return "Employers List", ["Company", "Industry", "Accreditation", "Registered"], rows


def _staff_vacancies(identity, date_from=None):
    query = Vacancy.query.filter(Vacancy.deleted_at.is_(None))
    if date_from:
        query = query.filter(Vacancy.created_at >= date_from)
    vacancies = query.order_by(Vacancy.created_at.desc()).limit(2000).all()
    rows = [[v.title, v.employer_company.company_name if v.employer_company else "", v.status, _date(v.created_at)] for v in vacancies]
    return "Vacancies List", ["Title", "Employer", "Status", "Posted"], rows


def _staff_program_applications(identity, date_from=None):
    query = ProgramApplication.query
    if date_from:
        query = query.filter(ProgramApplication.created_at >= date_from)
    apps = query.order_by(ProgramApplication.created_at.desc()).limit(2000).all()
    rows = [
        [a.program_type.upper(), a.jobseeker_profile.full_name if a.jobseeker_profile else "", a.status, _date(a.created_at)]
        for a in apps
    ]
    return "Program Applications (SPES/DILP/OWWA)", ["Program", "Jobseeker", "Status", "Date"], rows


def _admin_audit_trail(identity, date_from=None):
    query = AuditTrail.query
    if date_from:
        query = query.filter(AuditTrail.created_at >= date_from)
    entries = query.order_by(AuditTrail.created_at.desc()).limit(2000).all()
    rows = [
        [e.action, e.module, e.user_email or "", e.status, e.created_at.strftime("%Y-%m-%d %H:%M") if e.created_at else ""]
        for e in entries
    ]
    return "Audit Trail", ["Action", "Module", "User", "Status", "Timestamp"], rows


def _admin_dashboard_summary(identity, date_from=None):
    from services.dashboard_service import build_summary

    summary = build_summary()
    rows = [[key.replace("_", " ").title(), value] for key, value in summary.items()]
    return "System Dashboard Summary", ["Metric", "Value"], rows


REPORT_CATALOG = {
    "my_applications": {"roles": {"jobseeker"}, "builder": _my_applications},
    "my_referrals": {"roles": {"jobseeker"}, "builder": _my_referrals},
    "my_employment": {"roles": {"jobseeker"}, "builder": _my_employment},
    "my_vacancies": {"roles": {"employer"}, "builder": _my_vacancies},
    "my_applicants": {"roles": {"employer"}, "builder": _my_applicants},
    "my_hires": {"roles": {"employer"}, "builder": _my_hires},
    "jobseekers_list": {"roles": {"staff", "admin"}, "builder": _staff_jobseekers},
    "employers_list": {"roles": {"staff", "admin"}, "builder": _staff_employers},
    "vacancies_list": {"roles": {"staff", "admin"}, "builder": _staff_vacancies},
    "program_applications": {"roles": {"staff", "admin"}, "builder": _staff_program_applications},
    "audit_trail": {"roles": {"admin"}, "builder": _admin_audit_trail},
    "dashboard_summary": {"roles": {"admin"}, "builder": _admin_dashboard_summary},
}


def _build_docx(title: str, columns: list, rows: list) -> bytes:
    import docx

    document = docx.Document()
    document.add_heading(title, level=1)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = str(col)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _build_txt(title: str, columns: list, rows: list) -> bytes:
    lines = [title, "=" * len(title), "", "\t".join(str(c) for c in columns)]
    lines += ["\t".join("" if v is None else str(v) for v in row) for row in rows]
    return "\n".join(lines).encode("utf-8")


def generate_report_file(report_key: str, fmt: str, identity: str, role: str, date_from=None):
    """Returns (filename, mimetype, file_bytes), or None if report_key isn't in this
    role's whitelist — the caller must treat None as "not permitted", not "not found"."""
    entry = REPORT_CATALOG.get(report_key)
    if not entry or role not in entry["roles"]:
        return None

    title, columns, rows = entry["builder"](identity, date_from)
    timestamp = now_manila().strftime("%Y%m%d_%H%M%S")

    if fmt == "pdf":
        from services.pdf_service import generate_table_report

        pdf_bytes = generate_table_report(
            _esc(title), [_esc(c) for c in columns], [[_esc(v) for v in row] for row in rows],
            now_manila().strftime("%Y-%m-%d"),
        )
        return f"{report_key}_{timestamp}.pdf", "application/pdf", pdf_bytes

    if fmt == "docx":
        return (
            f"{report_key}_{timestamp}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _build_docx(title, columns, rows),
        )

    if fmt == "txt":
        return f"{report_key}_{timestamp}.txt", "text/plain", _build_txt(title, columns, rows)

    from services.excel_service import build_excel_report

    buf = build_excel_report(title, columns, rows)
    return (
        f"{report_key}_{timestamp}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        buf.getvalue(),
    )
